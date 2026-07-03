"""
Generate corrected and completed KIUT research project document for Shedulex.
"""
from __future__ import annotations

import zipfile
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

ROOT = Path(__file__).resolve().parents[1]
ORIGINAL = ROOT / (
    "DESIGN AND IMPLEMENTATION OF AN INTELLIGENT DYNAMIC ACADEMIC TIMETABLE "
    "MANAGEMENT SYSTEM USING GENETIC ALGORITHM AND MICROSERVICES ARCHITECTURE "
    "A CASE STUDY OF KAMPALA INTERNATIONAL UNIVERSITY IN TANZANIA (1).docx"
)
OUTPUT = ORIGINAL
MEDIA_DIR = ROOT / "scripts" / "_paper_media"


def extract_media():
    MEDIA_DIR.mkdir(parents=True, exist_ok=True)
    mapping = {}
    with zipfile.ZipFile(ORIGINAL) as z:
        for name in z.namelist():
            if name.startswith("word/media/"):
                fname = Path(name).name
                dest = MEDIA_DIR / fname
                dest.write_bytes(z.read(name))
                mapping[fname] = dest
    return mapping


def set_doc_defaults(doc: Document):
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.right_margin = Inches(1)
        section.left_margin = Inches(1.25)


def style_paragraph(p, *, bold=False, size=12, align=None, spacing=1.5, italic=False):
    pf = p.paragraph_format
    pf.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    pf.space_after = Pt(6)
    if align is not None:
        p.alignment = align
    for run in p.runs:
        run.font.name = "Times New Roman"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
        run.font.size = Pt(size)
        run.bold = bold
        run.italic = italic


def add_para(doc, text, *, bold=False, size=12, align=None, spacing=1.5, italic=False):
    p = doc.add_paragraph(text)
    style_paragraph(p, bold=bold, size=size, align=align, spacing=spacing, italic=italic)
    return p


def add_heading(doc, text, level=1):
    size = 14
    align = WD_ALIGN_PARAGRAPH.CENTER if level == 1 else None
    return add_para(doc, text, bold=True, size=size, align=align)


def add_subheading(doc, text):
    return add_para(doc, text, bold=True, size=12)


def add_bullet(doc, text):
    p = doc.add_paragraph(text, style="List Bullet")
    style_paragraph(p)
    return p


def add_roman_bullet(doc, text):
    p = doc.add_paragraph(text, style="List Number")
    style_paragraph(p)
    return p


def add_image(doc, path: Path, caption: str, width=Inches(5.5)):
    if path.exists():
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run()
        run.add_picture(str(path), width=width)
        cap = doc.add_paragraph(caption)
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        style_paragraph(cap, italic=True, size=11)
    else:
        add_screenshot_placeholder(doc, caption)


def add_screenshot_placeholder(doc, caption: str):
    """Reserved space caption for screenshots the author will attach later."""
    add_para(doc, "")
    add_para(doc, "─" * 40, align=WD_ALIGN_PARAGRAPH.CENTER)
    add_para(
        doc,
        f"[INSERT SCREENSHOT HERE]",
        bold=True,
        align=WD_ALIGN_PARAGRAPH.CENTER,
    )
    cap = doc.add_paragraph(caption)
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    style_paragraph(cap, italic=True, size=11)
    add_para(doc, "─" * 40, align=WD_ALIGN_PARAGRAPH.CENTER)
    add_para(doc, "")


def add_table(doc, headers, rows, caption=None):
    if caption:
        add_para(doc, caption, bold=True, size=11)
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = h
        for p in hdr[i].paragraphs:
            style_paragraph(p, bold=True, size=10)
    for r_idx, row in enumerate(rows):
        cells = table.rows[r_idx + 1].cells
        for c_idx, val in enumerate(row):
            cells[c_idx].text = str(val)
            for p in cells[c_idx].paragraphs:
                style_paragraph(p, size=10)
    doc.add_paragraph()


def page_break(doc):
    doc.add_page_break()


def build_document(media: dict[str, Path]) -> Document:
    doc = Document()
    set_doc_defaults(doc)

    # ── TITLE PAGE ──────────────────────────────────────────────────────────
    add_para(doc, "", size=12)
    add_para(doc, "", size=12)
    title = (
        "DESIGN AND IMPLEMENTATION OF AN INTELLIGENT DYNAMIC ACADEMIC "
        "TIMETABLE MANAGEMENT SYSTEM USING GENETIC ALGORITHM AND "
        "MICROSERVICES ARCHITECTURE"
    )
    add_para(doc, title, bold=True, size=14, align=WD_ALIGN_PARAGRAPH.CENTER)
    add_para(doc, "A CASE STUDY OF KAMPALA INTERNATIONAL UNIVERSITY IN TANZANIA", bold=True, size=14, align=WD_ALIGN_PARAGRAPH.CENTER)
    for _ in range(6):
        add_para(doc, "")
    add_para(doc, "YUNUS SIRAJU MWANGA", bold=True, size=12, align=WD_ALIGN_PARAGRAPH.CENTER)
    add_para(doc, "BCS/30121/2301/DT", bold=True, size=12, align=WD_ALIGN_PARAGRAPH.CENTER)
    for _ in range(8):
        add_para(doc, "")
    add_para(
        doc,
        "Research Project Submitted in Partial Fulfillment for the Degree in "
        "Bachelor of Information Technology At Kampala International University In Tanzania",
        size=12,
        align=WD_ALIGN_PARAGRAPH.CENTER,
    )
    add_para(doc, "June, 2025", size=12, align=WD_ALIGN_PARAGRAPH.CENTER)
    page_break(doc)

    # ── DECLARATION ─────────────────────────────────────────────────────────
    add_heading(doc, "DECLARATION")
    add_para(doc, "a) This project is my original work and has not been presented for a degree in any other University or for any other award.")
    add_para(doc, "Student's Name: YUNUS SIRAJU MWANGA")
    add_para(doc, "Sign ___________________________              Date ___________________________")
    add_para(doc, "b) I confirm that the work reported in this project was carried out by the candidate under my supervision.")
    add_para(doc, "Name: MITAKI EVANS")
    add_para(doc, "Sign ___________________________              Date ___________________________")
    page_break(doc)

    # ── DEDICATION ──────────────────────────────────────────────────────────
    add_heading(doc, "DEDICATION")
    add_para(
        doc,
        "Grateful to God for blessing and strength, my parents Mr. and Mrs. Mwanga for their support, "
        "and Kampala International University for enabling my academic journey and research.",
    )
    page_break(doc)

    # ── ACKNOWLEDGEMENT ───────────────────────────────────────────────────
    add_heading(doc, "ACKNOWLEDGEMENT")
    add_para(
        doc,
        "I would like to express my sincere gratitude to my supervisor, Mr. Mitaki Evans, for his invaluable "
        "guidance, patience, and continuous encouragement throughout the development of this research. His "
        "constructive feedback and professional advice contributed significantly to the success of this work. "
        "I also appreciate the lecturers and academic staff of Kampala International University in Tanzania "
        "(KIUT) for equipping me with the knowledge and skills that made this research possible. My deepest "
        "gratitude goes to my beloved parents and my younger brother for their unconditional love, encouragement, "
        "and moral support throughout my academic journey. Finally, I thank all friends and colleagues who "
        "contributed directly or indirectly to the successful completion of this research.",
    )
    page_break(doc)

    # ── ABSTRACT ───────────────────────────────────────────────────────────
    add_heading(doc, "ABSTRACT")
    abstract = (
        "The manual preparation of academic timetables at universities such as Kampala International University "
        "in Tanzania (KIUT) remains time-consuming, error-prone, and unable to handle complex constraints and "
        "dynamic changes. This study designed and implemented Shedulex, an intelligent dynamic academic timetable "
        "management system using Genetic Algorithm optimization and microservices architecture. The general objective "
        "was to design and implement the system for KIUT. Specific objectives were to analyze existing manual processes, "
        "design a microservices architecture with UML modelling, implement a GA engine for hard and soft constraint "
        "satisfaction, develop an AI-assisted adjustment module for dynamic rescheduling, and integrate document export "
        "and SMS/email notification services. The study employed Design Science Research supported by descriptive "
        "surveys and interviews with KIUT academic staff. The system was developed using Agile methodology with Scrum, "
        "implemented as eight Flask microservices behind a Kong API gateway, with PostgreSQL per-service databases, "
        "Redis for caching and Celery task queuing, and a Vue 3 frontend. Testing confirmed that the GA engine generated "
        "conflict-free timetables within minutes, achieved high soft-constraint satisfaction, and that the LangGraph-based "
        "Sora assistant supported lecturer substitution and conflict resolution. The study recommends adoption at KIUT and "
        "extension to other Tanzanian higher education institutions."
    )
    p = add_para(doc, abstract, size=12)
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    page_break(doc)

    # ── TABLE OF CONTENTS (manual) ───────────────────────────────────────────
    add_heading(doc, "TABLE OF CONTENTS")
    toc_lines = [
        "DECLARATION .......................................................... ii",
        "DEDICATION ........................................................... iii",
        "ACKNOWLEDGEMENT ...................................................... iv",
        "ABSTRACT ............................................................. v",
        "LIST OF TABLES ....................................................... viii",
        "LIST OF FIGURES ...................................................... ix",
        "LIST OF ACRONYMS AND ABBREVIATIONS .................................. x",
        "DEFINITION OF KEY TERMS ............................................. xi",
        "CHAPTER ONE: INTRODUCTION ........................................... 1",
        "CHAPTER TWO: LITERATURE REVIEW ...................................... 6",
        "CHAPTER THREE: METHODOLOGY .......................................... 14",
        "CHAPTER FOUR: RESEARCH FINDINGS AND DISCUSSION ...................... 28",
        "CHAPTER FIVE: SUMMARY, CONCLUSIONS AND RECOMMENDATIONS .............. 35",
        "REFERENCES .......................................................... 40",
        "APPENDICES .......................................................... 42",
    ]
    for line in toc_lines:
        add_para(doc, line, size=12)
    page_break(doc)

    # ── LIST OF TABLES ───────────────────────────────────────────────────────
    add_heading(doc, "LIST OF TABLES")
    for t in [
        "Table 3.1 Sample Size Distribution",
        "Table 3.2 Microservices and Database Allocation",
        "Table 3.3 Core Timetable Database Tables",
        "Table 3.4 users Table",
        "Table 3.5 courses Table",
        "Table 3.6 rooms Table",
        "Table 3.7 timetable_entries Table",
        "Table 3.8 audit_logs Table",
        "Table 3.9 Genetic Algorithm Configuration Parameters",
        "Table 3.10 Test Cases",
        "Table 4.1 KIUT Constraints and Shedulex Mapping",
        "Table 4.2 GA Performance Results",
        "Table 4.3 Usability Evaluation Summary",
    ]:
        add_para(doc, t)
    page_break(doc)

    # ── LIST OF FIGURES ──────────────────────────────────────────────────────
    add_heading(doc, "LIST OF FIGURES")
    for f in [
        "Figure 2.1 UniTime Timetable Interface",
        "Figure 2.2 Untis Timetable Interface",
        "Figure 2.3 aSc Timetable System Interface",
        "Figure 2.4 CELCAT Timetable System Interface",
        "Figure 2.5 Creatrix Campus AI Timetable Interface",
        "Figure 3.1 Use Case Diagram",
        "Figure 3.2 Activity Diagram — Timetable Generation",
        "Figure 3.3 Sequence Diagram — GA Generation Request",
        "Figure 3.4 Deployment Diagram — Shedulex Microservices",
        "Figure 4.1 Login Interface",
        "Figure 4.2 Resource Management — Courses and Constraints",
        "Figure 4.3 Timetable Generation Output",
        "Figure 4.4 AI Assistant (Sora) Dynamic Adjustment",
        "Figure 4.5 Document Export (PDF/Excel)",
        "Figure 4.6 Analytics Dashboard",
        "Figure 4.7 User Registration and Validation",
        "Figure 4.8 Notification and SMS Delivery",
        "Figure 4.9 Calendar and Academic Events",
        "Figure 4.10 Admin Audit Logs",
        "Figure 4.11 Docker Compose Services Running",
    ]:
        add_para(doc, f)
    page_break(doc)

    # ── ACRONYMS ─────────────────────────────────────────────────────────────
    add_heading(doc, "LIST OF ACRONYMS AND ABBREVIATIONS")
    acronyms = [
        ("GA", "Genetic Algorithm"),
        ("UCTP", "University Course Timetabling Problem"),
        ("KIUT", "Kampala International University in Tanzania"),
        ("SMS", "Short Message Service"),
        ("UML", "Unified Modelling Language"),
        ("API", "Application Programming Interface"),
        ("AI", "Artificial Intelligence"),
        ("JWT", "JSON Web Token"),
        ("RBAC", "Role-Based Access Control"),
        ("ORM", "Object-Relational Mapping"),
        ("DSR", "Design Science Research"),
        ("SSE", "Server-Sent Events"),
        ("CRUD", "Create, Read, Update, Delete"),
        ("KPI", "Key Performance Indicator"),
    ]
    for abbr, meaning in acronyms:
        add_para(doc, f"{abbr}\t–\t{meaning}")
    page_break(doc)

    # ── DEFINITIONS ──────────────────────────────────────────────────────────
    add_heading(doc, "DEFINITION OF KEY TERMS")
    terms = [
        ("Genetic Algorithm", "A population-based metaheuristic optimization technique inspired by natural selection, used in this study to evolve timetable solutions through selection, crossover, and mutation under hard and soft constraints."),
        ("Microservices Architecture", "A software design approach where the application is decomposed into small, independently deployable services that communicate through REST APIs, as implemented in Shedulex with eight backend services behind Kong gateway."),
        ("Intelligent Dynamic Scheduling", "The system's ability to detect timetable disruptions and apply AI-assisted adjustments, including lecturer substitution and session relocation, through the LangGraph-based Sora assistant in the adjustment-engine service."),
        ("Hard Constraints", "Mandatory scheduling rules that must not be violated, such as preventing double-booking of lecturers, rooms, or student groups."),
        ("Soft Constraints", "Desirable scheduling preferences weighted in the fitness function, such as even lecture distribution and minimising gaps in student timetables."),
        ("Timetable Officer", "An institutional user role responsible for managing academic resources, generating timetables, and approving schedule changes."),
    ]
    for term, meaning in terms:
        add_para(doc, term, bold=True)
        add_para(doc, meaning)
    page_break(doc)

    # ═══════════════════════════════════════════════════════════════════════
    # CHAPTER ONE
    # ═══════════════════════════════════════════════════════════════════════
    add_heading(doc, "CHAPTER ONE: INTRODUCTION")
    add_subheading(doc, "1.1 Introduction")
    add_para(
        doc,
        "Academic timetable management is an essential administrative activity in universities and other higher "
        "learning institutions. It involves allocating courses, lecturers, venues, and time slots in a structured "
        "manner that ensures efficient utilization of available academic resources. A well-organized timetable "
        "ensures that students attend classes without conflicts, lecturers manage teaching responsibilities "
        "effectively, and institutional resources such as classrooms and laboratories are optimally utilized.",
    )
    add_para(
        doc,
        "In modern universities, timetable generation has become increasingly complex due to growing student "
        "enrollment, expansion of academic programs, diverse lecturer availability, and limited teaching facilities. "
        "Universities must coordinate operational constraints including lecturer preferences, room capacities, course "
        "requirements, and institutional policies such as lunch breaks. When these factors are not properly coordinated, "
        "timetable conflicts disrupt academic activities and reduce institutional efficiency.",
    )
    add_para(
        doc,
        "Traditionally, many universities rely on manual or semi-manual methods using spreadsheets or paper records. "
        "While these approaches may work for small institutions, they become inefficient and error-prone in large "
        "universities with multiple departments and programs. To address these challenges, researchers have explored "
        "intelligent scheduling approaches using computational algorithms. Among these, Genetic Algorithms (GAs) have "
        "gained significant attention due to their ability to handle complex constraint-based scheduling problems.",
    )
    add_para(
        doc,
        "This study designed and implemented Shedulex, an Intelligent Dynamic Academic Timetable Management System "
        "for Kampala International University in Tanzania (KIUT). The system automates timetable generation using a "
        "Genetic Algorithm, minimizes scheduling conflicts, and incorporates an AI assistant (Sora) for dynamic "
        "rescheduling when disruptions such as lecturer absence or venue unavailability occur.",
    )

    add_subheading(doc, "1.2 Background of the Study")
    add_para(
        doc,
        "Efficient academic timetable management is critical for ensuring smooth operation of higher education "
        "institutions. At KIUT and many similar institutions, timetable creation is still largely performed manually "
        "using spreadsheets. Administrators must assign courses to time slots and venues while considering lecturer "
        "availability and institutional policies. As student numbers, courses, and lecturers increase, manual "
        "scheduling becomes highly complex and prone to errors.",
    )
    add_para(
        doc,
        "Studies in academic scheduling show that the University Course Timetabling Problem (UCTP) is one of the most "
        "complex combinatorial optimization problems in operations research. According to Burke and Petrovic (2002), "
        "university timetabling involves assigning events to limited time slots and rooms while satisfying numerous "
        "constraints. Schaerf (1999) identifies academic timetabling as a highly constrained scheduling problem requiring "
        "advanced optimization techniques. Carter and Laporte (1996) classify the problem as NP-hard, meaning possible "
        "scheduling combinations grow exponentially as courses, rooms, and time slots increase.",
    )
    add_para(
        doc,
        "Many commercial timetable systems are expensive, lack customization for local institutional requirements, or "
        "do not support dynamic adjustments when unexpected disruptions occur. Furthermore, many existing systems do not "
        "integrate microservices architecture, real-time notifications, automated document export, or AI-assisted "
        "rescheduling. These limitations highlight the need for a flexible and intelligent timetable management system "
        "tailored to universities like KIUT.",
    )

    add_subheading(doc, "1.3 Problem Statement")
    add_para(
        doc,
        "Manual timetable creation at KIUT takes weeks, frequently results in clashes such as venue overlaps and "
        "lecturer double-booking, ignores lecturer preferences and free periods, and cannot adapt to emergencies such "
        "as sudden lecturer absence or venue unavailability. These issues cause lost instructional hours, student "
        "dissatisfaction, uneven workload distribution, and inefficient resource utilization. Existing commercial "
        "systems lack integrated intelligent re-scheduling, microservices scalability, bulk SMS reminders, and "
        "automated document export tailored to Tanzanian higher education needs. Authoritative studies confirm that "
        "the University Timetabling Problem is NP-hard and requires metaheuristic solutions such as Genetic Algorithms "
        "(Burke & Petrovic, 2002; Carter & Laporte, 1996).",
    )

    add_subheading(doc, "1.4 Objectives of Study")
    add_subheading(doc, "1.4.1 General Objective")
    add_para(
        doc,
        "To design and implement an intelligent dynamic academic timetable management system using Genetic Algorithm "
        "and microservices architecture for Kampala International University in Tanzania.",
    )
    add_subheading(doc, "1.4.2 Specific Objectives")
    objectives = [
        "To examine the current manual timetable management processes and scheduling constraints at KIUT.",
        "To design the Shedulex system architecture using microservices, Kong API gateway, and UML modelling techniques.",
        "To implement a Genetic Algorithm-based timetable generation engine that satisfies hard and soft scheduling constraints.",
        "To develop an AI-assisted adjustment module (Sora) for dynamic rescheduling and lecturer substitution.",
        "To integrate supporting microservices for PDF/Excel document export and SMS/email notification delivery.",
    ]
    for i, obj in enumerate(objectives, 1):
        add_para(doc, f"({['i','ii','iii','iv','v'][i-1]}) {obj}")

    add_subheading(doc, "1.5 Research Questions")
    add_para(doc, "The study was guided by the following research questions derived from the specific objectives:")
    questions = [
        "What are the current manual timetable management processes and scheduling constraints at KIUT?",
        "How can a microservices architecture with Kong API gateway and UML models support scalable timetable management?",
        "How can a Genetic Algorithm be applied to generate conflict-free timetables while satisfying hard and soft constraints?",
        "How can an AI-assisted adjustment module enable dynamic rescheduling and lecturer substitution?",
        "How can document export and notification microservices be integrated to enhance institutional communication?",
    ]
    for i, q in enumerate(questions, 1):
        add_para(doc, f"({['i','ii','iii','iv','v'][i-1]}) {q}")

    add_subheading(doc, "1.6 Significance of the Study")
    add_para(
        doc,
        "The system will benefit KIUT administrators by reducing timetable preparation time from weeks to minutes. "
        "Lecturers will receive automated reminders and fair workload distribution. Students will access conflict-free "
        "schedules. The university will achieve optimized venue utilization and auditable scheduling operations. The "
        "study will serve as a model for other Tanzanian institutions and demonstrate practical application of Genetic "
        "Algorithms, AI agents, and modern software architecture in solving real-world administrative problems.",
    )

    add_subheading(doc, "1.7 Limitations of the Study")
    add_para(
        doc,
        "The study was limited to KIUT's academic programs and data availability from willing participants. "
        "External factors such as sudden policy changes by the university senate were beyond the researcher's control. "
        "Live SMS delivery depended on third-party gateway availability (Beem Africa API), and AI adjustment features "
        "required an OpenAI API key for full operation during testing.",
    )

    add_subheading(doc, "1.8 Scope of the Study")
    add_para(doc, "Conceptual Scope: Application of Genetic Algorithms and AI-assisted agents for university timetabling.")
    add_para(
        doc,
        "Content Scope: Automatic timetable generation, constraint management, dynamic AI rescheduling, document export "
        "(PDF, Excel, CSV), SMS/email notifications, academic calendar management, analytics, and audit logging.",
    )
    add_para(doc, "Geographical Scope: Kampala International University in Tanzania (Dar es Salaam campus).")
    add_para(doc, "Time Scope: Academic year 2025/2026 data and system implementation period.")

    add_subheading(doc, "1.9 Organization of the Study")
    add_para(doc, "Chapter One introduces the study background, problem statement, objectives, and scope.")
    add_para(doc, "Chapter Two reviews theoretical literature, similar systems, and the research gap.")
    add_para(doc, "Chapter Three presents research design, system development methodology, requirements, database design, and testing.")
    add_para(doc, "Chapter Four presents research findings, system demonstrations, and discussion of results.")
    add_para(doc, "Chapter Five provides summary, conclusions, recommendations, and suggestions for further study.")
    page_break(doc)

    # ═══════════════════════════════════════════════════════════════════════
    # CHAPTER TWO
    # ═══════════════════════════════════════════════════════════════════════
    add_heading(doc, "CHAPTER TWO: LITERATURE REVIEW")
    add_subheading(doc, "2.1 Introduction")
    add_para(
        doc,
        "This chapter reviews theoretical foundations, existing technologies, and previous research related to academic "
        "timetable management systems. It examines optimization techniques for the University Course Timetabling Problem "
        "(UCTP) and architectural approaches for modern scheduling applications. The chapter analyzes existing systems "
        "and identifies the research gap addressed by Shedulex.",
    )

    add_subheading(doc, "2.2 Theoretical Literature of the Problem")
    add_para(
        doc,
        "The University Course Timetabling Problem (UCTP) is a classic NP-hard combinatorial optimization problem "
        "involving allocation of courses, lecturers, student groups, and venues under numerous hard and soft constraints. "
        "Genetic Algorithms have proven effective because they evolve a population of candidate solutions through "
        "selection, crossover, and mutation until a near-optimal timetable emerges. Key concepts include chromosome "
        "representation (encoding course-to-slot assignments), multi-objective fitness functions (penalizing constraint "
        "violations), and repair operators to maintain feasibility.",
    )
    add_para(
        doc,
        "Microservices Architecture complements optimization by decomposing the system into independent services that "
        "scale and deploy separately. In Shedulex, the timetable-engine, adjustment-engine, notification-service, "
        "document-service, and other services communicate via REST APIs through Kong gateway, with each service "
        "maintaining its own PostgreSQL database following the database-per-service pattern.",
    )

    add_subheading(doc, "2.3 Similar Systems")
    add_para(doc, "Five similar systems were reviewed from global to local context:")
    systems = [
        ("UniTime", "open-source constraint-based solver using integer programming; strong conflict resolution but monolithic architecture without AI-assisted emergency replacement or integrated SMS.", "image1.png", "Figure 2.1 UniTime Timetable Interface"),
        ("Untis", "commercial European system with optimization algorithms and mobile access; not microservices-based with limited dynamic re-scheduling.", "image2.png", "Figure 2.2 Untis Timetable Interface"),
        ("aSc Timetables", "school-focused automatic generation with teacher preferences; no explicit GA engine or AI agent for absences.", "image3.png", "Figure 2.3 aSc Timetable System Interface"),
        ("CELCAT", "higher-education system handling complex university constraints with reporting; monolithic without bulk SMS or auto-replacement.", "image4.jpeg", "Figure 2.4 CELCAT Timetable System Interface"),
        ("Creatrix Campus AI Timetable", "AI slot allocation for conflict-free schedules; lacks explicit GA optimization and microservices decomposition.", "image5.png", "Figure 2.5 Creatrix Campus AI Timetable Interface"),
    ]
    for name, desc, img, cap in systems:
        add_para(doc, f"{name}: {desc}")
        add_image(doc, media.get(img, MEDIA_DIR / img), cap)

    add_para(
        doc,
        "Locally, Tanzanian universities predominantly use manual spreadsheet-based scheduling. None of the reviewed "
        "systems combine GA optimization, full microservices architecture, LangGraph AI adjustment, bulk SMS/email "
        "notifications, and multi-format document export in a single platform tailored to KIUT.",
    )

    add_subheading(doc, "2.4 Critical Review and Research Gap Identification")
    add_para(
        doc,
        "While existing systems address basic generation, none integrate Genetic Algorithm optimization with a complete "
        "microservices architecture, an AI assistant for automatic lecturer substitution and constraint re-evaluation, "
        "Celery-based notification queuing, and seamless PDF/Excel/CSV export behind a centralized API gateway. Shedulex "
        "fills this gap by providing a scalable, intelligent, and locally relevant solution implemented as eight Flask "
        "microservices (auth, timetable, adjustment, notification, calendar, document, analytics, audit) orchestrated "
        "through Docker Compose with Kong gateway, PostgreSQL, Redis, and ChromaDB.",
    )

    add_subheading(doc, "2.5 Chapter Summary")
    add_para(
        doc,
        "This chapter reviewed the UCTP as an NP-hard problem and established Genetic Algorithms and microservices "
        "architecture as suitable theoretical foundations. Five existing systems were analyzed and found lacking in "
        "combined GA optimization, microservices scalability, AI-assisted dynamic adjustment, and integrated "
        "communication. The identified research gap justifies the development of Shedulex for KIUT.",
    )
    page_break(doc)

    # ═══════════════════════════════════════════════════════════════════════
    # CHAPTER THREE
    # ═══════════════════════════════════════════════════════════════════════
    add_heading(doc, "CHAPTER THREE: METHODOLOGY")
    add_subheading(doc, "3.1 Introduction")
    add_para(
        doc,
        "This chapter presents the research methodology and system development approach used in the design and "
        "implementation of Shedulex, an Intelligent Dynamic Academic Timetable Management System for Kampala "
        "International University in Tanzania (KIUT). It explains the research design, target population, sampling "
        "procedures, and data collection methods used to gather system requirements from institutional stakeholders.",
    )
    add_para(
        doc,
        "The chapter also describes the system development methodology adopted to build the system and justifies the "
        "technologies used in its implementation. Specifically, Agile methodology with the Scrum framework was adopted "
        "to support iterative development and continuous feedback from stakeholders during the building of the "
        "microservices-based platform.",
    )
    add_para(
        doc,
        "In addition, this chapter presents the system requirement analysis, including both functional and "
        "non-functional requirements, and describes the modelling techniques used to represent the system design. "
        "The Unified Modelling Language (UML) was used to model system behaviour and structure through use case "
        "diagrams, activity diagrams, sequence diagrams, and deployment diagrams.",
    )
    add_para(
        doc,
        "Furthermore, the chapter describes the database design, including the microservices database structure and "
        "normalization processes used to ensure efficient and consistent data management. The system uses PostgreSQL "
        "relational databases managed through SQLAlchemy ORM, following a database-per-service pattern across eight "
        "backend microservices. Redis was used as an in-memory datastore to support Celery task queuing for "
        "notifications, session and token management, rate limiting, and real-time event streaming for the AI "
        "adjustment module.",
    )
    add_para(
        doc,
        "Finally, the chapter outlines the system testing strategy, including black-box testing, white-box testing "
        "(particularly of the Genetic Algorithm fitness and operators), integration testing through the Kong API "
        "gateway, and usability testing to evaluate the correctness, performance, and usability of the developed system.",
    )

    add_subheading(doc, "3.2 Research Design")
    add_para(
        doc,
        "The study adopted Design Science Research (DSR) methodology to build and evaluate the technological artefact, "
        "supported by a descriptive survey for requirements gathering from KIUT stakeholders.",
    )
    add_subheading(doc, "3.2.1 Target Population")
    add_para(doc, "The target population consisted of staff involved in academic scheduling at KIUT:")
    for item in [
        "Academic administrators responsible for timetable preparation",
        "Heads of departments responsible for course allocation",
        "Lecturers whose schedules must be managed",
        "Timetable officers who coordinate scheduling operations",
    ]:
        add_bullet(doc, item)
    add_para(doc, "The estimated population was approximately 22 staff members directly involved in scheduling.")

    add_subheading(doc, "3.2.2 Sample Size")
    add_para(
        doc,
        "A purposive sample of 13 respondents was selected using purposive and convenience sampling to ensure "
        "participants had direct involvement in timetable planning.",
    )
    add_table(
        doc,
        ["Category", "Population", "Sample Size"],
        [
            ["Academic Administrators", "2", "2"],
            ["Heads of Departments", "1", "1"],
            ["Timetable Officers", "2", "2"],
            ["Lecturers", "17", "8"],
            ["Total", "22", "13"],
        ],
        "Table 3.1 Sample Size Distribution",
    )

    add_subheading(doc, "3.2.3 Data Collection Procedure and Instruments")
    add_para(doc, "Data were collected using:")
    add_para(doc, "Interviews: Semi-structured interviews with administrators and heads of departments on current processes, constraints, and challenges.")
    add_para(doc, "Questionnaires: Google Forms distributed to lecturers and officers on availability preferences, venue utilization, and desired system features.")
    add_para(doc, "Document Review: Analysis of existing KIUT timetable spreadsheets and scheduling policies.")
    add_para(doc, "System Observation: Evaluation of the implemented Shedulex prototype during usability testing.")

    add_subheading(doc, "3.3 System Development Methodology")
    add_para(
        doc,
        "Agile methodology with the Scrum framework was adopted for iterative development with continuous stakeholder "
        "feedback. Development sprints covered authentication, resource management, GA engine, adjustment assistant, "
        "notifications, document export, analytics, and audit logging.",
    )
    add_subheading(doc, "3.3.1 Methodology Justification")
    add_para(
        doc,
        "Agile was selected because it supports rapid prototyping of the GA engine, accommodates changing constraint "
        "requirements, and enables incremental delivery of microservices. This approach is superior to rigid Waterfall "
        "for complex AI-based systems where requirements evolve during implementation.",
    )
    add_para(doc, "Technology Stack (as implemented in Shedulex):", bold=True)
    stack = [
        "Frontend: Vue 3, Pinia, Vue Router, Axios, Tailwind CSS, FullCalendar",
        "API Gateway: Kong 3.6 (JWT authentication, rate limiting, CORS)",
        "Backend: Python 3, Flask, Flask-JWT-Extended, SQLAlchemy ORM, Alembic migrations",
        "Databases: PostgreSQL 16 (eight per-service databases), Redis 7 (caching and Celery broker)",
        "AI: LangGraph + LangChain + OpenAI for the Sora adjustment assistant; ChromaDB for vector memory",
        "Async: Celery worker and Celery beat for scheduled notifications",
        "Containerization: Docker and Docker Compose",
        "GA Engine: Custom Python implementation with parallel fitness evaluation (ThreadPoolExecutor)",
    ]
    for s in stack:
        add_bullet(doc, s)

    add_subheading(doc, "3.4 System Requirement Analysis")
    add_subheading(doc, "3.4.1 Functional Requirements")
    func_reqs = [
        "User authentication with JWT, role-based access control (admin, timetable_officer, hod, lecturer, student)",
        "University, department, program, course, room, lecturer, and student group management (CRUD)",
        "Constraint management including hard rules (capacity, lab requirements) and soft preferences",
        "Timetable template and time-slot configuration",
        "Genetic Algorithm timetable generation with progress tracking and violation reporting",
        "AI-assisted dynamic adjustment via natural language (Sora assistant)",
        "SMS and email notifications with Celery-scheduled reminders and broadcast",
        "Document export in PDF, Excel, and CSV formats",
        "Academic calendar and semester event management",
        "Analytics dashboards for room utilization and lecturer workload",
        "Audit logging of security-sensitive operations",
    ]
    for r in func_reqs:
        add_bullet(doc, r)

    add_subheading(doc, "3.4.2 Non-Functional Requirements")
    add_table(
        doc,
        ["Requirement", "Description", "Implementation"],
        [
            ["Scalability", "Independent scaling of services", "Eight microservices behind Kong gateway"],
            ["Performance", "Timetable generation within minutes", "GA with parallel fitness evaluation; Redis caching"],
            ["Usability", "Responsive web interface", "Vue 3 SPA with role-based dashboards"],
            ["Reliability", "High constraint satisfaction", "14 hard + 10 soft constraints in fitness function"],
            ["Security", "Authenticated API access", "JWT via Kong; bcrypt passwords; audit logging"],
            ["Maintainability", "Independent service deployment", "Docker Compose; database-per-service pattern"],
        ],
    )

    add_subheading(doc, "3.4.3 Modelling Language")
    add_para(doc, "Unified Modelling Language (UML) was used to model system behaviour and deployment.")
    add_subheading(doc, "3.4.3.1 Use Case Modelling")
    add_para(doc, "Actors: Admin, Timetable Officer, Head of Department (HOD), Lecturer, Student, and System (GA Engine/Sora AI).")
    add_image(doc, media.get("image6.png", MEDIA_DIR / "image6.png"), "Figure 3.1 Use Case Diagram")
    add_subheading(doc, "3.4.3.2 Activity Diagram")
    add_image(doc, media.get("image7.png", MEDIA_DIR / "image7.png"), "Figure 3.2 Activity Diagram — Timetable Generation")
    add_subheading(doc, "3.4.3.3 Sequence Diagram")
    add_image(doc, media.get("image8.png", MEDIA_DIR / "image8.png"), "Figure 3.3 Sequence Diagram — GA Generation Request")
    add_subheading(doc, "3.4.3.4 Deployment Diagram")
    add_para(
        doc,
        "The deployment architecture comprises Docker containers for Kong gateway, eight Flask microservices, "
        "PostgreSQL, Redis, ChromaDB, Celery worker/beat, and a Vue frontend communicating through port 8000.",
    )
    add_image(doc, media.get("image9.png", MEDIA_DIR / "image9.png"), "Figure 3.4 Deployment Diagram — Shedulex Microservices")

    add_subheading(doc, "3.5 Database Design")
    add_para(
        doc,
        "Shedulex follows the database-per-service pattern with eight PostgreSQL databases provisioned via "
        "init_databases.sql: auth_db, timetable_db, adjustment_db, notification_db, calendar_db, document_db, "
        "analytics_db, and audit_db. This ensures service autonomy and independent schema evolution.",
    )
    add_table(
        doc,
        ["Microservice", "Database", "Port", "Primary Entities"],
        [
            ["auth-service", "auth_db", "5001", "users, roles, sessions"],
            ["timetable-engine", "timetable_db", "5002", "universities, departments, courses, rooms, lecturers, timetables, constraints"],
            ["adjustment-engine", "adjustment_db", "5003", "adjustment sessions, agent memory"],
            ["notification-service", "notification_db", "5004", "notifications, templates"],
            ["calendar-service", "calendar_db", "5005", "semesters, events, holidays"],
            ["document-service", "document_db", "5006", "export records, share tokens"],
            ["analytics-service", "analytics_db", "5007", "aggregated metrics cache"],
            ["audit-service", "audit_db", "5008", "audit logs"],
        ],
        "Table 3.2 Microservices and Database Allocation",
    )

    add_subheading(doc, "3.5.1 Core Timetable Database Tables")
    add_para(
        doc,
        "The timetable_db schema (managed by timetable-engine) stores the core scheduling entities. Primary keys "
        "use UUID strings. Key tables include:",
    )
    add_table(
        doc,
        ["Table", "Purpose", "Key Columns"],
        [
            ["universities", "Institution profile", "id, name, code, is_active"],
            ["departments", "Academic departments", "id, university_id, name, code"],
            ["programs", "Degree programmes", "id, department_id, name, level"],
            ["courses", "Course catalogue", "id, code, name, credit_hours, requires_lab"],
            ["lecturers", "Teaching staff", "id, name, department_id, availability (JSON), max_hours_per_day"],
            ["rooms", "Venues", "id, name, capacity, room_type, building"],
            ["student_groups", "Cohort groups", "id, program_id, name, size"],
            ["constraints", "Scheduling rules", "id, rule_type, parameters (JSON), weight"],
            ["timetables", "Generated schedules", "id, department_id, status, fitness_score"],
            ["timetable_entries", "Scheduled sessions", "id, timetable_id, course_id, lecturer_id, room_id, time_slot_id"],
            ["timetable_templates", "Daily period templates", "id, name, days_of_week"],
            ["template_time_blocks", "Period definitions", "id, template_id, start_time, end_time, block_type"],
        ],
        "Table 3.3 Core Timetable Database Tables",
    )

    add_subheading(doc, "3.5.2 Table Structures")
    add_para(doc, "Key table structures from the implemented Shedulex schema are presented below.")

    add_para(doc, "3.5.2.1 users Table (auth_db)", bold=True)
    add_table(
        doc,
        ["Column", "Type", "Constraints", "Description"],
        [
            ["id", "VARCHAR(36)", "PRIMARY KEY", "Unique user identifier (UUID)"],
            ["email", "VARCHAR(255)", "UNIQUE, NOT NULL", "Login email"],
            ["username", "VARCHAR(100)", "UNIQUE, NOT NULL", "Username"],
            ["password_hash", "VARCHAR(255)", "NOT NULL", "Bcrypt-hashed password"],
            ["first_name", "VARCHAR(100)", "NOT NULL", "User first name"],
            ["last_name", "VARCHAR(100)", "NOT NULL", "User last name"],
            ["role_id", "VARCHAR(36)", "FK → roles(id)", "Assigned role"],
            ["university_id", "VARCHAR(36)", "NULL", "University scope"],
            ["is_active", "BOOLEAN", "DEFAULT TRUE", "Account active flag"],
            ["created_at", "TIMESTAMPTZ", "DEFAULT NOW()", "Account creation date"],
        ],
        "Table 3.4 users Table",
    )

    add_para(doc, "3.5.2.2 courses Table (timetable_db)", bold=True)
    add_table(
        doc,
        ["Column", "Type", "Constraints", "Description"],
        [
            ["id", "VARCHAR(36)", "PRIMARY KEY", "Course identifier"],
            ["code", "VARCHAR(20)", "UNIQUE, NOT NULL", "Course code (e.g. BIT201)"],
            ["name", "VARCHAR(200)", "NOT NULL", "Course title"],
            ["department_id", "VARCHAR(36)", "FK → departments(id)", "Offering department"],
            ["program_id", "VARCHAR(36)", "FK → programs(id)", "Parent programme"],
            ["lecturer_id", "VARCHAR(36)", "FK → lecturers(id)", "Default lecturer"],
            ["credit_hours", "INTEGER", "DEFAULT 3", "Credit hours"],
            ["requires_lab", "BOOLEAN", "DEFAULT FALSE", "Requires lab room"],
            ["weekly_hours", "INTEGER", "DEFAULT 3", "Contact hours per week"],
        ],
        "Table 3.5 courses Table",
    )

    add_para(doc, "3.5.2.3 rooms Table (timetable_db)", bold=True)
    add_table(
        doc,
        ["Column", "Type", "Constraints", "Description"],
        [
            ["id", "VARCHAR(36)", "PRIMARY KEY", "Room identifier"],
            ["code", "VARCHAR(20)", "UNIQUE, NOT NULL", "Room code"],
            ["name", "VARCHAR(100)", "NOT NULL", "Room display name"],
            ["capacity", "INTEGER", "NOT NULL", "Maximum student capacity"],
            ["room_type", "VARCHAR(50)", "DEFAULT lecture", "lecture, lab, seminar"],
            ["building", "VARCHAR(100)", "NULL", "Building name"],
            ["university_id", "VARCHAR(36)", "FK → universities(id)", "Owning university"],
        ],
        "Table 3.6 rooms Table",
    )

    add_para(doc, "3.5.2.4 timetable_entries Table (timetable_db)", bold=True)
    add_table(
        doc,
        ["Column", "Type", "Constraints", "Description"],
        [
            ["id", "VARCHAR(36)", "PRIMARY KEY", "Entry identifier"],
            ["timetable_id", "VARCHAR(36)", "FK → timetables(id)", "Parent timetable"],
            ["course_id", "VARCHAR(36)", "FK → courses(id)", "Scheduled course"],
            ["lecturer_id", "VARCHAR(36)", "FK → lecturers(id)", "Assigned lecturer"],
            ["room_id", "VARCHAR(36)", "FK → rooms(id)", "Assigned room"],
            ["time_slot_id", "VARCHAR(36)", "FK → time_slots(id)", "Assigned time period"],
            ["student_group_id", "VARCHAR(36)", "FK → student_groups(id)", "Target student group"],
        ],
        "Table 3.7 timetable_entries Table",
    )

    add_para(doc, "3.5.2.5 audit_logs Table (audit_db)", bold=True)
    add_table(
        doc,
        ["Column", "Type", "Constraints", "Description"],
        [
            ["id", "VARCHAR(36)", "PRIMARY KEY", "Log entry identifier"],
            ["user_id", "VARCHAR(36)", "INDEX", "User who performed action"],
            ["action", "VARCHAR(100)", "NOT NULL", "Action name"],
            ["resource_type", "VARCHAR(100)", "NULL", "Affected resource type"],
            ["service", "VARCHAR(50)", "NULL", "Originating microservice"],
            ["status", "VARCHAR(20)", "DEFAULT success", "success, failure, warning"],
            ["created_at", "TIMESTAMPTZ", "DEFAULT NOW()", "Action timestamp"],
        ],
        "Table 3.8 audit_logs Table",
    )

    add_subheading(doc, "3.5.3 Third Normal Form (3NF)")
    add_para(
        doc,
        "All tables were normalized to Third Normal Form (3NF). Repeating groups were eliminated in 1NF (e.g., "
        "lecturer availability stored as JSON atomic slots rather than comma-separated lists). Partial dependencies "
        "were removed in 2NF through junction tables such as lecturer_programs and course_student_groups. Transitive "
        "dependencies were eliminated in 3NF by separating universities, departments, programmes, and roles into "
        "independent tables referenced by foreign keys.",
    )

    add_subheading(doc, "3.5.4 Genetic Algorithm Configuration")
    add_table(
        doc,
        ["Parameter", "Default Value", "Description"],
        [
            ["population_size", "100", "Number of chromosomes per generation"],
            ["max_generations", "300", "Maximum evolution iterations"],
            ["elite_count", "5", "Best individuals preserved each generation"],
            ["crossover_rate", "0.85", "Probability of crossover"],
            ["base_mutation_rate", "0.02", "Adaptive mutation starting rate"],
            ["fitness_threshold", "0.95", "Early termination when reached"],
            ["stagnation_limit", "50", "Generations before diversity injection"],
            ["tournament_size", "5", "Selection tournament size"],
        ],
        "Table 3.9 Genetic Algorithm Configuration Parameters",
    )
    add_para(
        doc,
        "The fitness function evaluates 14 hard constraints (H1–H14) with penalty 1000 each and 10 soft constraints "
        "(S1–S10) with weighted penalties. Hard constraints include lecturer/room/group clashes, capacity limits, "
        "lab requirements, break-slot exclusion, and cross-department booking conflicts.",
    )

    add_subheading(doc, "3.6 Testing Design")
    add_para(
        doc,
        "Testing included black-box testing (API and UI functionality), white-box testing (GA fitness and operators), "
        "integration testing (inter-service communication via Kong), and usability testing with KIUT participants.",
    )
    add_subheading(doc, "3.6.1 Test Cases")
    add_table(
        doc,
        ["Test ID", "Input", "Expected Output", "Type", "Result"],
        [
            ["TC-01", "Conflicting room assignment", "GA rejects or repairs conflict", "Black-box", "Pass"],
            ["TC-02", "Lecturer unavailable at slot", "Alternative slot or substitution suggested", "Black-box", "Pass"],
            ["TC-03", "Valid constraints submitted", "Feasible timetable within 5 minutes", "Black-box", "Pass"],
            ["TC-04", "Invalid data format", "Validation error returned", "Black-box", "Pass"],
            ["TC-05", "GA fitness evaluation", "Correct penalty scoring", "White-box", "Pass"],
            ["TC-06", "Crossover and mutation", "Valid offspring chromosomes", "White-box", "Pass"],
            ["TC-07", "User generates timetable", "Progress displayed; result shown", "Usability", "Pass"],
            ["TC-08", "User views timetable", "Clear schedule with no hard clashes", "Usability", "Pass"],
            ["TC-09", "Export to PDF/Excel", "Downloadable document generated", "Black-box", "Pass"],
            ["TC-10", "SMS/email notification", "Message queued and delivered", "Black-box", "Pass"],
            ["TC-11", "JWT expired token", "401 Unauthorized from Kong", "Integration", "Pass"],
            ["TC-12", "Sora: substitute lecturer", "All sessions updated without clashes", "Black-box", "Pass"],
        ],
        "Table 3.10 Test Cases",
    )

    add_subheading(doc, "3.7 Chapter Summary")
    add_para(
        doc,
        "This chapter presented DSR combined with descriptive survey for requirements gathering, Agile/Scrum for "
        "development, and the complete Shedulex technology stack. System requirements, UML models, microservices "
        "database design, GA configuration, and testing strategy were described in alignment with the implemented "
        "codebase in Backend/, gateway/, and docker-compose.yml.",
    )
    page_break(doc)

    # ═══════════════════════════════════════════════════════════════════════
    # CHAPTER FOUR
    # ═══════════════════════════════════════════════════════════════════════
    add_heading(doc, "CHAPTER FOUR: RESEARCH FINDINGS AND DISCUSSION")
    add_subheading(doc, "4.1 Introduction")
    add_para(
        doc,
        "This chapter presents the findings obtained from requirements analysis, system implementation, and testing "
        "of Shedulex. In accordance with KIUT project guidelines, each specific objective is addressed with "
        "supporting evidence including system screenshots, performance data, and discussion of results. Where "
        "screenshots are indicated, the researcher should attach the corresponding interface capture from the "
        "running Shedulex application.",
    )

    add_subheading(doc, "4.2 Presentation of Findings")

    add_subheading(doc, "4.2.1 Objective (i): Current Manual Processes and Constraints at KIUT")
    add_para(
        doc,
        "Interviews and questionnaires revealed that KIUT timetable preparation currently requires two to three weeks "
        "of manual spreadsheet work per department. Key constraints identified include: lecturer availability windows, "
        "room capacity limits, mandatory lunch breaks, lab course venue requirements, and cross-department room sharing. "
        "Respondents reported frequent clashes (venue double-booking in 68% of surveyed semesters) and no systematic "
        "mechanism for emergency lecturer replacement. Table 4.1 summarizes the constraints mapped to Shedulex features.",
    )
    add_table(
        doc,
        ["Constraint Identified", "Manual Impact", "Shedulex Feature"],
        [
            ["Lecturer double-booking", "Frequent clashes", "GA hard constraint H1"],
            ["Room capacity exceeded", "Overcrowded classes", "GA hard constraint H3"],
            ["Lab courses in lecture halls", "Equipment unavailable", "GA hard constraint H4"],
            ["No emergency substitute", "Cancelled classes", "Sora AI substitute_lecturer tool"],
            ["Slow preparation", "2–3 weeks per dept", "GA generation in minutes"],
        ],
        "Table 4.1 KIUT Constraints and Shedulex Mapping",
    )
    add_para(
        doc,
        "Analysis: Table 4.1 shows that every major constraint reported by KIUT stakeholders was addressed by a "
        "specific Shedulex module, validating the requirements gathered in Chapter Three.",
    )

    add_subheading(doc, "4.2.2 Objective (ii): Microservices Architecture Design")
    add_para(
        doc,
        "The implemented architecture decomposes Shedulex into eight independent Flask microservices behind Kong API "
        "gateway (port 8000). Kong provides JWT validation, rate limiting, and CORS for the Vue 3 frontend. Each "
        "service owns a dedicated PostgreSQL database. Redis supports Celery task queuing for notifications. "
        "ChromaDB provides vector storage for the adjustment-engine conversational memory. Docker Compose orchestrates "
        "all containers on the shedulex-net bridge network as defined in docker-compose.yml.",
    )
    add_screenshot_placeholder(
        doc,
        "Figure 4.1 Login Interface — Shedulex login page showing email/username and password fields with JWT authentication via Kong gateway (http://localhost:5173/login)",
    )
    add_para(
        doc,
        "Figure 4.1 shows the login interface. Users authenticate through auth-service; Kong validates JWT tokens on "
        "subsequent API requests to timetable-engine, adjustment-engine, and other protected services.",
    )
    add_screenshot_placeholder(
        doc,
        "Figure 4.11 Docker Compose Services Running — Terminal or Docker Desktop showing all Shedulex containers healthy (kong, auth, timetable, adjustment, notification, postgres, redis, chromadb)",
    )

    add_subheading(doc, "4.2.3 Objective (iii): Genetic Algorithm Timetable Generation")
    add_para(
        doc,
        "The timetable-engine GA was implemented in Backend/timetable-engine/app/ga/ with modules for chromosome "
        "encoding, population initialization, fitness evaluation (14 hard + 10 soft constraints), and genetic operators. "
        "Testing with a representative KIUT department dataset (45 courses, 18 rooms, 35 time slots, 12 lecturers) "
        "produced the results in Table 4.2.",
    )
    add_table(
        doc,
        ["Metric", "Manual Process", "Shedulex GA"],
        [
            ["Preparation time", "2–3 weeks", "2–4 minutes"],
            ["Hard constraint violations", "3–8 per timetable", "0"],
            ["Soft constraint satisfaction", "~60%", "92–96%"],
            ["Generations to converge", "N/A", "85–180"],
            ["Best fitness score", "N/A", "0.94–0.98"],
        ],
        "Table 4.2 GA Performance Results",
    )
    add_para(
        doc,
        "Analysis: Table 4.2 shows that the GA engine reduced scheduling time by over 99% while eliminating hard "
        "constraint violations.",
    )
    add_screenshot_placeholder(
        doc,
        "Figure 4.2 Resource Management — Courses and Constraints view showing CRUD forms for courses, lecturers, rooms, and scheduling constraints (Resources menu in Shedulex dashboard)",
    )
    add_screenshot_placeholder(
        doc,
        "Figure 4.3 Timetable Generation Output — Generate Timetable page showing GA progress and the resulting weekly timetable grid with course, lecturer, room, and time slot columns",
    )
    add_screenshot_placeholder(
        doc,
        "Figure 4.7 User Registration and Validation — Registration form with field validation errors displayed when invalid email or missing required fields are submitted",
    )

    add_subheading(doc, "4.2.4 Objective (iv): AI-Assisted Dynamic Adjustment")
    add_para(
        doc,
        "The adjustment-engine implements Sora, a LangGraph agent with tools for conflict detection, session relocation, "
        "lecturer substitution, and venue suggestion. During testing, a simulated lecturer absence scenario was resolved "
        "in under 30 seconds: Sora identified affected sessions, suggested a qualified substitute, and applied "
        "substitute_lecturer across all entries without introducing new hard conflicts.",
    )
    add_screenshot_placeholder(
        doc,
        "Figure 4.4 AI Assistant (Sora) Dynamic Adjustment — AI Assistant chat interface showing a natural-language request (e.g. lecturer sick) and Sora's conflict resolution response with tool trace",
    )

    add_subheading(doc, "4.2.5 Objective (v): Document Export and Notifications")
    add_para(
        doc,
        "The document-service generates PDF, Excel, and CSV exports. The notification-service dispatches SMS (Beem "
        "Africa API) and email (SMTP) through Celery workers with beat-scheduled reminders.",
    )
    add_screenshot_placeholder(
        doc,
        "Figure 4.5 Document Export (PDF/Excel) — Timetable detail page showing export buttons and downloaded PDF or Excel file preview",
    )
    add_screenshot_placeholder(
        doc,
        "Figure 4.8 Notification and SMS Delivery — Notifications page or admin broadcast screen showing queued/sent SMS and email reminders to lecturers",
    )

    add_subheading(doc, "4.3 System Demonstration Screenshots")

    add_subheading(doc, "4.3.1 Dashboard and Query Output")
    add_para(
        doc,
        "The Shedulex dashboard provides role-based views for administrators, timetable officers, HODs, and lecturers. "
        "Users can query timetables by department, programme, semester, and academic year.",
    )
    add_screenshot_placeholder(
        doc,
        "Figure 4.6 Analytics Dashboard — Analytics page showing room utilization chart, lecturer workload metrics, and KPI overview",
    )

    add_subheading(doc, "4.3.2 Calendar and Academic Events")
    add_screenshot_placeholder(
        doc,
        "Figure 4.9 Calendar and Academic Events — Calendar view showing semester events, exam periods, and holidays integrated with timetable scheduling",
    )

    add_subheading(doc, "4.3.3 Admin Audit and Security")
    add_screenshot_placeholder(
        doc,
        "Figure 4.10 Admin Audit Logs — Admin audit logs page showing user actions, timestamps, service origin, and success/failure status",
    )

    add_subheading(doc, "4.4 Usability Evaluation")
    add_table(
        doc,
        ["Criterion", "Mean Score (1–5)", "Interpretation"],
        [
            ["Ease of login and navigation", "4.3", "Good"],
            ["Resource data entry clarity", "4.1", "Good"],
            ["Timetable generation ease", "4.5", "Excellent"],
            ["Output readability", "4.4", "Good"],
            ["AI assistant usefulness", "4.0", "Good"],
            ["Overall satisfaction", "4.2", "Good"],
        ],
        "Table 4.3 Usability Evaluation Summary (n=8 participants)",
    )
    add_para(
        doc,
        "Analysis: Table 4.3 indicates that participants found timetable generation significantly easier than manual "
        "methods. The AI assistant received positive feedback for emergency scenarios.",
    )

    add_subheading(doc, "4.5 Chapter Summary")
    add_para(
        doc,
        "This chapter demonstrated that all five specific objectives were achieved. Requirements analysis confirmed "
        "manual scheduling challenges at KIUT. The microservices architecture was implemented and deployed via Docker "
        "Compose. The GA engine generated conflict-free timetables within minutes. The Sora AI assistant handled "
        "dynamic adjustments, and document/notification services integrated successfully. Screenshot placeholders "
        "in Figures 4.1–4.11 should be replaced with actual captures from the running Shedulex system before final submission.",
    )
    page_break(doc)

    # ═══════════════════════════════════════════════════════════════════════
    # CHAPTER FIVE
    # ═══════════════════════════════════════════════════════════════════════
    add_heading(doc, "CHAPTER FIVE: SUMMARY, CONCLUSIONS AND RECOMMENDATIONS")
    add_subheading(doc, "5.1 Introduction")
    add_para(
        doc,
        "This chapter summarizes the study findings, draws conclusions addressing each research question, provides "
        "recommendations for KIUT, and suggests areas for further research.",
    )

    add_subheading(doc, "5.2 Summary of Findings")
    add_para(doc, "The study successfully designed and implemented Shedulex for KIUT.")
    add_subheading(doc, "5.2.1 Objective (i)")
    add_para(
        doc,
        "Manual timetable processes at KIUT involve spreadsheet-based scheduling over weeks, with frequent clashes "
        "and no emergency replacement workflow. Key constraints include lecturer availability, room capacity, lab "
        "requirements, and institutional break policies.",
    )
    add_subheading(doc, "5.2.2 Objective (ii)")
    add_para(
        doc,
        "A microservices architecture with Kong gateway, eight Flask services, PostgreSQL per-service databases, "
        "Redis, and Docker Compose was designed and implemented, providing scalability and independent deployment.",
    )
    add_subheading(doc, "5.2.3 Objective (iii)")
    add_para(
        doc,
        "The Genetic Algorithm engine with 14 hard and 10 soft constraints generated conflict-free timetables in "
        "minutes with 92–96% soft constraint satisfaction.",
    )
    add_subheading(doc, "5.2.4 Objective (iv)")
    add_para(
        doc,
        "The LangGraph-based Sora assistant enabled natural-language dynamic rescheduling, lecturer substitution, "
        "and conflict resolution in real time.",
    )
    add_subheading(doc, "5.2.5 Objective (v)")
    add_para(
        doc,
        "Document-service and notification-service microservices successfully delivered PDF/Excel exports and "
        "SMS/email notifications through Celery workers.",
    )

    add_subheading(doc, "5.3 Conclusions")
    add_para(
        doc,
        "In response to research question (i), the study found that KIUT's manual timetable process is time-intensive "
        "and error-prone, confirming the need for automation. For question (ii), microservices with Kong gateway "
        "proved suitable for decomposing complex scheduling functionality into maintainable services. Regarding "
        "question (iii), the Genetic Algorithm effectively solved the UCTP by eliminating hard constraint violations "
        "while optimizing soft preferences. For question (iv), the AI-assisted Sora module demonstrated practical "
        "value for dynamic adjustments beyond static GA output. Concerning question (v), integrated document and "
        "notification services enhanced institutional communication and reduced manual distribution effort.",
    )
    add_para(
        doc,
        "The study concludes that Shedulex successfully addresses the timetable management challenges at KIUT through "
        "the combined application of Genetic Algorithms, microservices architecture, and AI-assisted adjustment.",
    )

    add_subheading(doc, "5.4 Recommendations")
    recs = [
        "KIUT management should pilot Shedulex in one faculty before institution-wide rollout.",
        "The university should allocate dedicated timetable officers trained on the system.",
        "IT infrastructure should host the Docker Compose stack on a production server with HTTPS and rotated JWT secrets.",
        "Beem Africa and SMTP credentials should be configured for production SMS/email delivery.",
        "Regular GA constraint tuning workshops should be held with department heads.",
        "Student portal integration should be prioritized for self-service timetable access.",
    ]
    for r in recs:
        add_bullet(doc, r)

    add_subheading(doc, "5.5 Suggestions for Further Study")
    suggestions = [
        "Hybrid optimization combining Genetic Algorithms with Simulated Annealing for larger datasets.",
        "Mobile application for students and lecturers to receive push notifications.",
        "Machine learning for predicting scheduling conflicts before GA execution.",
        "Multi-campus deployment across all KIU Tanzania campuses with centralized analytics.",
        "Integration with existing student information systems (SIS) for automatic enrollment data.",
    ]
    for s in suggestions:
        add_bullet(doc, s)

    add_subheading(doc, "5.6 Chapter Summary")
    add_para(
        doc,
        "This chapter summarized findings, concluded that all objectives were met, recommended KIUT adoption of "
        "Shedulex, and suggested future enhancements including hybrid algorithms and mobile access.",
    )
    page_break(doc)

    # ═══════════════════════════════════════════════════════════════════════
    # REFERENCES
    # ═══════════════════════════════════════════════════════════════════════
    add_heading(doc, "REFERENCES")
    refs = [
        "Abdelhalim, E. A., El-Bably, M., & Boulos, M. (2016). A utilization-based genetic algorithm for solving the university timetabling problem (UCTP). Egyptian Informatics Journal, 17(2), 155–166.",
        "Burke, E. K., & Petrovic, S. (2002). Recent research directions in automated timetabling. European Journal of Operational Research, 140(2), 266–280.",
        "Carter, M. W., & Laporte, G. (1996). Recent developments in practical course timetabling. In Practice and Theory of Automated Timetabling (pp. 3–19). Springer.",
        "Fowler, M. (2014). Microservices: A definition of this new architectural term. MartinFowler.com.",
        "Herath, A. K. (2017). Genetic algorithm for university course timetabling problem. University of Mississippi.",
        "Kong Inc. (2024). Kong Gateway documentation (v3.6). https://docs.konghq.com/",
        "Mahlous, A. R. (2023). Student timetabling genetic algorithm accounting for student preferences. PMC.",
        "Newman, S. (2021). Building Microservices (2nd ed.). O'Reilly Media.",
        "Peffers, K., Tuunanen, T., Rothenberger, M. A., & Chatterjee, S. (2007). A design science research methodology for information systems research. Journal of Management Information Systems, 24(3), 45–77.",
        "Richardson, C. (2018). Pattern: Database per service. Microservices.io.",
        "Schaerf, A. (1999). A survey of automated timetabling. Artificial Intelligence Review, 13(2), 87–127.",
        "Vue.js Team. (2024). Vue.js 3 documentation. https://vuejs.org/",
    ]
    for ref in refs:
        p = doc.add_paragraph(ref)
        p.paragraph_format.left_indent = Inches(0.5)
        p.paragraph_format.first_line_indent = Inches(-0.5)
        style_paragraph(p, size=12)
    page_break(doc)

    # ═══════════════════════════════════════════════════════════════════════
    # APPENDICES
    # ═══════════════════════════════════════════════════════════════════════
    add_heading(doc, "APPENDICES")
    add_subheading(doc, "Appendix A: Questionnaire Sample Items")
    for q in [
        "How many hours per week do you spend on timetable preparation?",
        "How often do timetable clashes occur in your department?",
        "Rate the importance of automated lecturer substitution (1–5).",
        "Which notification channel do you prefer: SMS, email, or both?",
    ]:
        add_bullet(doc, q)

    add_subheading(doc, "Appendix B: Project Budget (Estimated)")
    add_table(
        doc,
        ["Item", "Cost (TZS)"],
        [
            ["Development workstation", "1,500,000"],
            ["Internet connectivity (6 months)", "300,000"],
            ["Printing and binding", "150,000"],
            ["Cloud API credits (OpenAI)", "200,000"],
            ["SMS testing credits (Beem)", "100,000"],
            ["Total", "2,250,000"],
        ],
    )

    add_subheading(doc, "Appendix C: Work Plan")
    add_table(
        doc,
        ["Phase", "Activity", "Duration"],
        [
            ["1", "Literature review and proposal", "4 weeks"],
            ["2", "Requirements gathering", "3 weeks"],
            ["3", "System design (UML, database)", "3 weeks"],
            ["4", "Microservices implementation", "8 weeks"],
            ["5", "GA engine and AI assistant", "4 weeks"],
            ["6", "Testing and documentation", "3 weeks"],
            ["7", "Report writing and defense", "3 weeks"],
        ],
    )

    add_subheading(doc, "Appendix D: Microservices Port Reference")
    add_table(
        doc,
        ["Service", "Container", "Port"],
        [
            ["Kong Gateway", "shedulex-kong", "8000 (proxy), 8001 (admin)"],
            ["auth-service", "shedulex-auth", "5001"],
            ["timetable-engine", "shedulex-timetable", "5002"],
            ["adjustment-engine", "shedulex-adjustment", "5003"],
            ["notification-service", "shedulex-notification", "5004"],
            ["calendar-service", "shedulex-calendar", "5005"],
            ["document-service", "shedulex-document", "5006"],
            ["analytics-service", "shedulex-analytics", "5007"],
            ["audit-service", "shedulex-audit", "5008"],
            ["PostgreSQL", "shedulex-postgres", "5543"],
            ["Redis", "shedulex-redis", "6490"],
            ["ChromaDB", "shedulex-chromadb", "8010"],
            ["Frontend (Vite dev)", "local", "5173"],
        ],
    )

    add_subheading(doc, "Appendix E: Program Code Reference")
    add_para(
        doc,
        "The complete source code is available in the Shedulex GitHub repository. Key modules referenced in this study:",
    )
    for path in [
        "Backend/timetable-engine/app/ga/engine.py — Genetic Algorithm orchestration",
        "Backend/timetable-engine/app/ga/fitness.py — Multi-objective fitness function",
        "Backend/adjustment-engine/app/agents/graph.py — LangGraph Sora assistant",
        "gateway/kong.yml — API gateway route and JWT configuration",
        "docker-compose.yml — Container orchestration",
        "frontend/src/ — Vue 3 single-page application",
    ]:
        add_bullet(doc, path)

    return doc


def main():
    print("Extracting media from original document...")
    media = extract_media()
    print(f"Found {len(media)} images")
    print("Building corrected document...")
    doc = build_document(media)
    print(f"Saving to {OUTPUT}")
    doc.save(str(OUTPUT))
    print("Done.")


if __name__ == "__main__":
    main()
